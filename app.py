from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from sentence_transformers import util
import pickle
import docx
import PyPDF2
import re
from temp import get_jobs_and_skills   # Import your job-fetching API module
import pandas as pd
import numpy as np

import asyncio
import nest_asyncio
nest_asyncio.apply()

import streamlit as st

# Load pre-trained models and other resources at the start
svc_model = pickle.load(open('clf.pkl', 'rb'))
tfidf = pickle.load(open('tfidf.pkl', 'rb'))
le = pickle.load(open('encoder.pkl', 'rb'))


from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
# Load BERT Model
bert_model = SentenceTransformer('all-MiniLM-L6-v2')  # Lightweight & fast
model = SentenceTransformer('paraphrase-mpnet-base-v2')  # upgrade

import spacy
nlp = spacy.load("en_core_web_sm")

def split_sentences(text):
    doc = nlp(text)
    return [sent.text.strip() for sent in doc.sents if sent.text.strip()]

# Custom CSS for styling
st.markdown(
    """
    <style>
    .center-header {
        text-align: center;
        font-size: 2.5rem;
        font-weight: bold;
        color: #4CAF50;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Header
st.markdown('<h1 class="center-header">RESUME ANALYZER AND JOB FINDER</h1>', unsafe_allow_html=True)

# Main options using Streamlit Radio
option = st.radio("Choose what you want to do:", ["Find Jobs Based on Skills", "Analyze Resume"])


import os
import requests

def download_model():
    file_id = "16O889rgUvMWA09vWd04fn90uBIxCOLpf"
    url = f"https://drive.google.com/uc?export=download&id={file_id}"
    response = requests.get(url)
    with open("clf.pkl", "wb") as f:
        f.write(response.content)

# Download only if not already present
if not os.path.exists("clf.pkl"):
    download_model()

try:
    asyncio.get_running_loop()
except RuntimeError:
    asyncio.set_event_loop(asyncio.new_event_loop())
    
def get_Jd_Array(array, clean_jd, predefined_skills):
    # Traverse through the clean_jd and add matching predefined skills to the array
    for skill in predefined_skills:
        if skill.lower() in clean_jd.lower():
            array.append(skill)
    return array

def get_Resume_Score(clean_resume, jd):
    # Count how many skills from jd are in the clean_resume
    score = 0
    for skill in jd:
        if skill.lower() in clean_resume.lower():
            score += 1
    return score

    
# Define helper functions
def cleanResume(txt):
    """Clean resume text by removing unwanted characters."""
    cleanText = re.sub(r'http\S+', ' ', txt)  # Remove URLs
    cleanText = re.sub(r'RT|cc', ' ', cleanText)  # Remove 'RT' and 'cc'
    cleanText = re.sub(r'#\S+', ' ', cleanText)  # Remove hashtags
    cleanText = re.sub(r'@\S+', ' ', cleanText)  # Remove mentions
    cleanText = re.sub(r'[!"#$%&\'()*+,-./:;<=>?@\[\\\]^_{|}~]+', ' ', cleanText)  # Remove punctuations
    cleanText = re.sub(r'[^\x00-\x7f]', ' ', cleanText)  # Remove non-ASCII characters
    cleanText = re.sub(r'\s+', ' ', cleanText).strip()  # Remove extra whitespace
    return cleanText



def calculate_bert_similarity(resume_text, jd_text):
    """
    Convert resume & job description to BERT embeddings and compute similarity.
    """
    resume_embedding = bert_model.encode([resume_text])
    jd_embedding = bert_model.encode([jd_text])

    similarity = cosine_similarity(resume_embedding, jd_embedding)[0][0]
    return round(similarity * 100, 2)  # Convert to percentage

def extract_keywords(text):
    # Basic keyword extractor using regex
    words = re.findall(r'\b\w+\b', text.lower())
    keywords = set([word for word in words if word not in ENGLISH_STOP_WORDS and len(word) > 2])
    return keywords

def sentence_level_match(resume_text, jd_text, model, top_k=10):
    # 1. Use spaCy to split sentences
    resume_sents = split_sentences(resume_text)
    jd_sents = split_sentences(jd_text)

    # 2. Encode all sentences
    resume_emb = model.encode(resume_sents, convert_to_tensor=True)
    jd_emb = model.encode(jd_sents, convert_to_tensor=True)

    # 3. Cosine similarity matrix
    cosims = util.pytorch_cos_sim(resume_emb, jd_emb).cpu().numpy()

    # 4. Find best matches
    matches = []
    for i, r in enumerate(resume_sents):
        j_idx = cosims[i].argmax()  # Best matching JD sentence index
        score = round(float(cosims[i][j_idx] * 100), 2)  # Convert to percentage
        matches.append((r, jd_sents[j_idx], score))

    # 5. Sort and pick top_k matches
    top_matches = sorted(matches, key=lambda x: -x[2])[:top_k]

    # 6. Convert to a DataFrame
    df_matches = pd.DataFrame(top_matches, columns=['Resume Sentence', 'JD Sentence', 'Similarity (%)'])

    return df_matches

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ''
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + ' '
    return text.strip()



def extract_text_from_docx(file):
    """Extract text from DOCX files."""
    doc = docx.Document(file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text


def extract_text_from_txt(file):
    """Extract text from TXT files."""
    try:
        text = file.read().decode('utf-8')
    except UnicodeDecodeError:
        text = file.read().decode('latin-1')  # Fallback encoding
    return text


def pred(input_resume):
    """Predict the category of a resume."""
    cleaned_text = cleanResume(input_resume)
    vectorized_text = tfidf.transform([cleaned_text]).toarray()
    predicted_category = svc_model.predict(vectorized_text)
    return le.inverse_transform(predicted_category)[0]

# Read the CSV file
df = pd.read_csv("UpdatedResumeDataSet.csv")

# Extract the 2nd column (index 1) into a NumPy array
second_column_array = df.iloc[:, 1].to_numpy()

# Split elements on spaces, hyphens, or underscores; remove empty strings and strip white spaces
split_words = [
    word.strip() for sentence in second_column_array 
    for word in re.split(r'[ \-_]', str(sentence)) 
    if word.strip()  # Remove null or empty strings
]

# Remove duplicates using a tuple
unique_words = tuple(set(split_words))

# Convert back to a NumPy array
predefined_skills = np.array(unique_words)
    
if option == "Find Jobs Based on Skills":
    st.subheader("Job Finder")
    st.title("Resume and Job Finder App")
    st.markdown("Upload a resume in PDF, TXT, or DOCX format, get the predicted job category, and find relevant jobs.")

    # File upload section
    uploaded_file = st.file_uploader("Upload a Resume", type=["pdf", "docx", "txt"])

    if uploaded_file is not None:
        try:
            # Extract text from the uploaded file
            file_extension = uploaded_file.name.split('.')[-1].lower()
            if file_extension == 'pdf':
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                resume_text = ''.join(page.extract_text() for page in pdf_reader.pages)
            elif file_extension == 'docx':
                doc = docx.Document(uploaded_file)
                resume_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            elif file_extension == 'txt':
                resume_text = uploaded_file.read().decode('utf-8')
            else:
                raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")

            # Clean resume text
            cleaned_resume_text = cleanResume(resume_text)

            # Display extracted text (optional)
            if st.checkbox("Show extracted text", False):
                st.text_area("Extracted Resume Text", resume_text, height=300)

            # Make prediction
            st.subheader("Predicted Category")
            category = pred(cleaned_resume_text)
            st.markdown(
                f"""
                <div style="
                    display: flex;
                    justify-content: center;
                    align-items: center;
                    height: 150px;
                    color: #4CAF50;
                    font-size: 3rem;
                    font-weight: bold;
                    text-align: center;
                    border: 2px solid #4CAF50;
                    border-radius: 10px;
                    padding: 20px;
                    background-color: #f9f9f9;">
                    The predicted category is: {category}
                </div>
                """,
                unsafe_allow_html=True,
            )
            st.write(f"The predicted category of the uploaded resume is: **{category}**")

            # Fetch job recommendations
            st.subheader("Job Recommendations with Matching Scores")
            city = st.text_input("Enter the city of interest:")
            if city and category:
                jobs_data = get_jobs_and_skills(city, category)

                if "error" in jobs_data[0]:
                    st.write(jobs_data[0]["error"])
                else:
                    # Calculate matching scores for each job
                    scored_jobs = []
                    for job in jobs_data:
                        job_description = job.get("description", "")
                        job_title = job.get("title", "")
                        combined_job_text = job_title + " " + job_description

                        d_array = []
                        jd_skills_score = get_Jd_Array(d_array, combined_job_text, predefined_skills)

                        # Calculate skills match for Resume
                        resume_skills_score = get_Resume_Score(cleaned_resume_text, d_array)
                        
                        similarity_score = (resume_skills_score / len(d_array)) * 100
                        similarity_score = round(similarity_score, 2)

                        # similarity_score = calculate_cosine_similarity(cleaned_resume_text, combined_job_text)

                        # Add job with score to list
                        scored_jobs.append({
                            "score": similarity_score,
                            "title": job.get("title"),
                            "company": job.get("company"),
                            "location": job.get("location"),
                            "employment_type": job.get("employment_type"),
                            "description": job.get("description"),
                            "url": job.get("url"),
                            "skills": job.get("skills")
                        })

                    # Sort jobs by similarity score in descending order
                    sorted_jobs = sorted(scored_jobs, key=lambda x: x["score"], reverse=True)

                    # Display sorted jobs
                    for job in sorted_jobs:
                        if job['score'] < 30:
                            color = "red"
                        elif job['score'] < 60:
                            color = "yellow"
                        elif job['score'] < 75:
                            color = "lightgreen"
                        else:
                            color = "green"
                        st.markdown(
                            f"""
                            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 20px; padding: 10px; border: 1px solid #ddd; border-radius: 10px;">
                                <div style="flex: 1; margin-right: 20px;">
                                    <p><strong>Job Title:</strong> {job['title']}</p>
                                    <p><strong>Company:</strong> {job['company']}</p>
                                    <p><strong>Location:</strong> {job['location']}</p>
                                    <p><strong>Job Type:</strong> {job['employment_type']}</p>
                                    <p><strong>Description:</strong> {job['description'][:150]}...</p>
                                    <p><strong>Skills Found:</strong> {', '.join(job['skills']) if job['skills'] else 'No tech skills found'}</p>
                                    <p><a href="{job['url']}" target="_blank"><strong>Apply Here</strong></a></p>
                                </div>
                                <div style="flex-shrink: 0; text-align: center; margin-left: 20px;">
                                    <p style="font-size: 2rem; font-weight: bold; color: {color};">Matching Score: {job['score']}%</p>
                                </div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )

        except Exception as e:
            st.error(f"Error processing the file: {str(e)}")



elif option == "Analyze Resume":
    st.subheader("Upload your Resume and Job Description")
    
    # File upload for Resume
    uploaded_resume = st.file_uploader("Upload Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
    uploaded_jd = st.text_area("Enter Job Description (or upload JD file below)", "")

    # Upload JD file if needed
    uploaded_jd_file = st.file_uploader("Or upload Job Description File", type=["pdf", "docx", "txt"])

    if uploaded_resume is not None and (uploaded_jd or uploaded_jd_file is not None):
        try:
            # Extract resume text
            if uploaded_resume is not None:
                resume_extension = uploaded_resume.name.split('.')[-1].lower()
                if resume_extension == 'pdf':
                    resume_text = extract_text_from_pdf(uploaded_resume)
                elif resume_extension == 'docx':
                    resume_text = extract_text_from_docx(uploaded_resume)
                elif resume_extension == 'txt':
                    resume_text = extract_text_from_txt(uploaded_resume)
                else:
                    st.error("Unsupported resume file format.")
            
            # Extract Job Description text
            if uploaded_jd_file is not None:
                jd_extension = uploaded_jd_file.name.split('.')[-1].lower()
                if jd_extension == 'pdf':
                    jd_text = extract_text_from_pdf(uploaded_jd_file)
                elif jd_extension == 'docx':
                    jd_text = extract_text_from_docx(uploaded_jd_file)
                elif jd_extension == 'txt':
                    jd_text = extract_text_from_txt(uploaded_jd_file)
                else:
                    st.error("Unsupported job description file format.")
            elif uploaded_jd:
                jd_text = uploaded_jd
            else:
                st.error("Please provide a job description.")

            # Clean the extracted text
            clean_resume = cleanResume(resume_text)
            clean_jd = cleanResume(jd_text)


            d_array = []
            jd_skills_score = get_Jd_Array(d_array, clean_jd, predefined_skills)

            # Calculate skills match for Resume
            resume_skills_score = get_Resume_Score(clean_resume, d_array)
            if len(d_array) == 0:
                st.error("No skills found in Job Description.")
            else:
                if resume_skills_score == 0:
                    st.error("No matching skills found in Resume.")
                else:
                    # Calculate match percentage based on how many skills in the JD are found in the resume
                    matching_score = calculate_bert_similarity(clean_resume, clean_jd)
                    
                    # Ensure the score does not exceed 100%
                    matching_score = min(matching_score, 100)

                st.markdown(
                    f"""
                    <div style="
                        display: flex;
                        justify-content: center;
                        align-items: center;
                        height: 200px;
                        color: #4CAF50;
                        font-size: 3rem;
                        font-weight: bold;
                        text-align: center;
                        border: 2px solid #4CAF50;
                        border-radius: 10px;
                        padding: 20px;
                        background-color: #f0f0f0;
                        ">
                        Matching Score: {matching_score:.2f}%
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                st.subheader("🧠 Detailed Sentence-Level Matching")
                placeholder = st.empty()
                placeholder.text("Performing sentence-level matching...")
                result = sentence_level_match(resume_text, jd_text, model, top_k=10)
                placeholder.text("Sentence-level matching completed!")
                placeholder.empty()  # Clear the placeholder
                st.dataframe(result)
                st.markdown("---")

        except Exception as e:
            st.error(f"Error: {str(e)}")
